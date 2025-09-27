import os
import logging
import pytesseract
import google.generativeai as genai
from PIL import Image
from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
import fitz  # PyMuPDF
from datetime import datetime
import cv2
import numpy as np
import traceback
import pandas as pd
import shutil
import threading
import time

# Logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.FileHandler("app.log"), logging.StreamHandler()],
)
logger = logging.getLogger(__name__)

# Gemini API
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Flask setup
app = Flask(__name__, template_folder="templates")
UPLOAD_FOLDER = "uploads"
ARCHIVE_FOLDER = "archive"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(ARCHIVE_FOLDER, exist_ok=True)

# Global per gestire timeout
processing_status = {}

# --- Configurazione Tesseract ---
def configure_tesseract():
    try:
        pytesseract.get_tesseract_version()
        logger.info("Tesseract configurato correttamente")
    except Exception as e:
        logger.error(f"Errore Tesseract: {e}")
        for path in ["/usr/bin/tesseract", "/usr/local/bin/tesseract"]:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                logger.info(f"Configurato Tesseract da path: {path}")
                break
configure_tesseract()

# --- Preprocessing immagini OTTIMIZZATO ---
def preprocess_image_for_ocr(image_path, fast_mode=True):
    """Preprocessing ottimizzato con modalità veloce per Render"""
    try:
        img = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)
        if img is None:
            return None
        
        if fast_mode:
            # Modalità veloce: solo resize e threshold semplice
            h, w = img.shape
            if max(h, w) > 1500:  # Ridimensiona se troppo grande
                scale = 1500 / max(h, w)
                img = cv2.resize(img, (int(w*scale), int(h*scale)))
            
            _, img = cv2.threshold(img, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        else:
            # Modalità completa
            img = cv2.adaptiveThreshold(img, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                        cv2.THRESH_BINARY, 31, 2)
            img = cv2.medianBlur(img, 3)
        
        return img
    except Exception as e:
        logger.error(f"Errore preprocessing: {e}")
        return None

def ocr_image_file(image_path, lang="ita+eng"):
    """
    OCR su immagine con preprocessing (usa preprocess_image_for_ocr se disponibile).
    Restituisce stringa (vuota se errore).
    """
    try:
        # usa la funzione di preprocessing già presente nel file
        processed = preprocess_image_for_ocr(image_path, fast_mode=True)
        if processed is not None:
            # processed è un numpy array (grayscale) -> converti in PIL
            pil_img = Image.fromarray(processed)
        else:
            pil_img = Image.open(image_path)
        text = pytesseract.image_to_string(pil_img, lang=lang)
        return text or ""
    except Exception as e:
        logger.error(f"OCR image error: {e}")
        return ""



# --- Estrazione testo OTTIMIZZATA ---
def extract_text_from_file_async(file_path, task_id):
    """
    Estrae testo da singolo file (PDF testuale o PDF scansione -> OCR pagina per pagina).
    Aggiorna processing_status[task_id] con {'status','progress','result'}.
    """
    processing_status[task_id] = {"status": "processing", "progress": 0, "start_time": time.time()}
    try:
        ext = os.path.splitext(file_path)[1].lower()
        result_text = ""

        if ext == ".pdf":
            # apri PDF con PyMuPDF (fitz)
            doc = fitz.open(file_path)
            total_pages = len(doc)
            for i, page in enumerate(doc):
                # aggiorna progress parziale
                processing_status[task_id]["progress"] = int((i / max(total_pages,1)) * 60)

                # prova estrazione testo digitale
                try:
                    page_text = page.get_text("text") or ""
                except Exception:
                    page_text = ""

                if page_text.strip():
                    result_text += f"--- Pagina {i+1} ---\n{page_text}\n\n"
                else:
                    # pagina immagine: rendi immagine ad alta risoluzione e OCR
                    pix = page.get_pixmap(dpi=300)
                    temp_img = os.path.join(UPLOAD_FOLDER, f"ocr_{os.path.basename(file_path)}_p{i+1}.png")
                    pix.save(temp_img)
                    ocr_t = ocr_image_file(temp_img)
                    result_text += f"--- Pagina {i+1} (OCR) ---\n{ocr_t}\n\n"
                    # pulizia immagine temporanea
                    try:
                        os.remove(temp_img)
                    except Exception:
                        pass

                processing_status[task_id]["progress"] = int(((i + 1) / max(total_pages,1)) * 90)

            doc.close()

        elif ext in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
            # immagine singola -> OCR
            result_text = ocr_image_file(file_path)
            processing_status[task_id]["progress"] = 100

        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                result_text = f.read()
            processing_status[task_id]["progress"] = 100

        else:
            result_text = f"[Formato non supportato: {ext}]"
            processing_status[task_id]["progress"] = 100

        processing_status[task_id]["status"] = "completed"
        processing_status[task_id]["result"] = result_text

    except Exception as e:
        logger.error(traceback.format_exc())
        processing_status[task_id]["status"] = "error"
        processing_status[task_id]["result"] = str(e)


# --- Prompt generator ---
def get_prompt(base_type, custom, text):
    # Limita il testo per evitare timeout API
    max_chars = 8000
    if len(text) > max_chars:
        text = text[:max_chars] + "\n...[testo troncato per performance]"
    
    if base_type == "simple":
        return f"Analizza questo referto medico e fornisci un riassunto chiaro e semplice per un paziente (max 300 parole):\n\n{text}"
    elif base_type == "intermediate":
        return f"""Analizza questo referto medico e fornisci un riassunto strutturato (max 400 parole):
        - Diagnosi principale
        - Parametri fuori norma con valori numerici
        - Terapie o raccomandazioni
        Testo referto:
        {text}"""
    elif base_type == "detailed":
        return f"""Analisi tecnica del referto medico (max 500 parole):
        - Diagnosi e classificazioni mediche
        - Valori di laboratorio con range normali
        - Correlazioni cliniche
        Testo referto:
        {text}"""
    elif base_type == "custom" and custom:
        return f"{custom}\n\nTesto referto:\n{text}"
    return f"Riassumi il seguente referto medico:\n{text}"

# --- Summarization OTTIMIZZATA ---
def generate_summary(prompt):
    try:
        model = genai.GenerativeModel("gemini-2.5-flash-lite-preview-09-2025")
        # Timeout più basso per Gemini
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        logger.error(f"Errore Gemini: {traceback.format_exc()}")
        return f"⚠️ Servizio temporaneamente non disponibile. Riprova tra qualche minuto.\nErrore: {str(e)}"

# --- Word export ---
def create_word_doc(summary, full_text):
    doc = Document()
    doc.add_heading("Riassunto Referto Medico", 0)
    doc.add_paragraph(summary)
    doc.add_page_break()
    doc.add_heading("Testo Integrale", level=1)
    doc.add_paragraph(full_text)
    file_path = os.path.join(ARCHIVE_FOLDER, "riassunto_referto.docx")
    doc.save(file_path)
    return file_path

# --- Routes OTTIMIZZATE ---
@app.route("/")
def home():
    return render_template("index.html")

@app.route("/upload", methods=["POST"])
def upload_file():
    try:
        task_id = str(int(time.time() * 1000))  # Timestamp come ID
        processing_status[task_id] = {"status": "starting", "start_time": time.time()}
        
        files = request.files.getlist("file")
        if not files or not files[0].filename:
            return jsonify({"error": "Nessun file caricato"}), 400
            
        # Controllo dimensioni
        total_size = sum(file.content_length or 0 for file in files if file.content_length)
        if total_size > 10 * 1024 * 1024:  # 10MB limit
            return jsonify({"error": "File troppo grandi. Limite: 10MB totali"}), 400
        
        # Salva file e avvia elaborazione asincrona
        filepaths = []
        for file in files:
            if file.filename:
                filepath = os.path.join(UPLOAD_FOLDER, f"{task_id}_{file.filename}")
                file.save(filepath)
                filepaths.append(filepath)
        
        # Avvia thread di elaborazione
        thread = threading.Thread(target=extract_text_from_files_thread, args=(filepaths, task_id))
        thread.daemon = True
        thread.start()
        
        return jsonify({"task_id": task_id, "status": "processing"})
        
    except Exception as e:
        logger.error(traceback.format_exc())
        return jsonify({"error": f"Errore upload: {str(e)}"}), 500

def extract_text_from_files_thread(filepaths, task_id):
    """
    Thread di coordinamento: elabora ogni file con extract_text_from_file_async e
    aggrega i risultati nel task principale task_id.
    """
    try:
        texts = []
        total_files = len(filepaths)
        for idx, filepath in enumerate(filepaths):
            file_task_id = f"{task_id}_file_{idx}"
            # esegui sincronamente l'estrazione (funzione aggiorna processing_status[file_task_id])
            extract_text_from_file_async(filepath, file_task_id)

            # recupera risultato (già impostato dalla funzione)
            res_info = processing_status.get(file_task_id, {})
            texts.append(res_info.get("result", ""))

            # rimuovi file temporaneo
            try:
                os.remove(filepath)
            except Exception:
                logger.warning(f"Impossibile rimuovere file temporaneo: {filepath}")

            # aggiorna progresso complessivo
            processing_status[task_id]["progress"] = int(((idx + 1) / max(total_files,1)) * 100)

        full_text = "\n\n".join(texts)
        processing_status[task_id]["status"] = "completed"
        processing_status[task_id]["result"] = full_text

    except Exception as e:
        logger.error(traceback.format_exc())
        processing_status[task_id]["status"] = "error"
        processing_status[task_id]["result"] = str(e)

@app.route("/check_status/<task_id>", methods=["GET"])
def check_status(task_id):
    """Controlla status elaborazione"""
    if task_id not in processing_status:
        return jsonify({"error": "Task non trovato"}), 404
    
    return jsonify(processing_status[task_id])

@app.route("/analyze", methods=["POST"])
def analyze_text():
    try:
        full_text = request.form.get("extracted_text", "").strip()
        if not full_text:
            return jsonify({"error": "Nessun testo da analizzare"}), 400
            
        prompt_type = request.form.get("prompt_type", "simple")
        custom_prompt = request.form.get("custom_prompt", "")
        
        prompt = get_prompt(prompt_type, custom_prompt, full_text)
        summary = generate_summary(prompt)
        
        # Salva solo se riuscito
        if not summary.startswith("⚠️"):
            create_word_doc(summary, full_text)
        
        return jsonify({"summary": summary, "status": "success"})
        
    except Exception as e:
        logger.error(traceback.format_exc())
        return jsonify({"error": f"Errore analisi: {str(e)}"}), 500

@app.route("/download-summary")
def download_summary():
    file_path = os.path.join(ARCHIVE_FOLDER, "riassunto_referto.docx")
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    return jsonify({"error": "File non disponibile"}), 404

@app.route("/reset", methods=["POST"])
def reset():
    try:
        # Pulisci cartelle e status
        for folder in [UPLOAD_FOLDER, ARCHIVE_FOLDER]:
            if os.path.exists(folder):
                shutil.rmtree(folder)
                os.makedirs(folder)
        
        processing_status.clear()
        logger.info("Reset completato")
        return jsonify({"status": "reset_done"})
        
    except Exception as e:
        logger.error(traceback.format_exc())
        return jsonify({"error": str(e)}), 500

@app.route("/health")
def health():
    return jsonify({"status": "OK", "timestamp": str(datetime.now())})

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    # Configurazione ottimizzata per Render
    app.run(host="0.0.0.0", port=port, debug=False, threaded=True)
