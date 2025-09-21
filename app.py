import os
import pytesseract
#import openai
import google.generativeai as genai # Aggiungi l'import di Gemini
#import gemini
from PIL import Image
from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
import fitz  # PyMuPDF
from dotenv import load_dotenv
#import google.generativeai as genai

# Specify the full path to the Tesseract executable
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'


# Carica la chiave API dal file .env
load_dotenv()
#openai.api_key = os.getenv("OPENAI_API_KEY")
# Configura l'API di Gemini con la chiave caricata
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Configurazione di Flask
app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
ARCHIVE_FOLDER = 'archive'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(ARCHIVE_FOLDER, exist_ok=True)

# Funzione per estrarre testo da un'immagine o PDF
def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    if ext == '.pdf':
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()

    elif ext in ['.png', '.jpg', '.jpeg']:
        text = pytesseract.image_to_string(Image.open(file_path), lang='ita')

    elif ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()

    elif ext == '.docx':
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + '\n'

    elif ext == '.xlsx':
        import pandas as pd
        xls = pd.read_excel(file_path, sheet_name=None)
        for sheet_name, df in xls.items():
            text += f"\n--- Foglio: {sheet_name} ---\n"
            text += df.to_string(index=False)

    else:
        text = "Formato non supportato."

    return text

# Funzione per generare il riassunto con Gemini
def generate_summary(text):
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(
            f"Riassumi questo referto medico in un linguaggio semplice. Evidenzia i valori fuori dalla norma, la diagnosi principale e le raccomandazioni del medico. Sii conciso. Il referto è:\n\n{text}"
        )
        return response.text
    except Exception as e:
        return f"Errore nella generazione del riassunto: {e}"


# Funzione per creare il file Word
def create_word_doc(summary, full_text):
    doc = Document()
    doc.add_heading('Riassunto Referto Medico', 0)
    doc.add_paragraph(summary)
    doc.add_page_break()
    doc.add_heading('Testo Integrale', level=1)
    doc.add_paragraph(full_text)
    
    file_path = os.path.join(ARCHIVE_FOLDER, "riassunto_referto.docx")
    doc.save(file_path)
    return file_path

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "Nessun file selezionato"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "Nessun file selezionato"}), 400
    
    filename = file.filename
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # Estrazione e analisi
    full_text = extract_text_from_file(filepath)
    simple_summary = generate_summary(full_text)
    
    # Crea e salva il file Word
    word_path = create_word_doc(simple_summary, full_text)

    return jsonify({
        "summary": simple_summary,
        "full_text": full_text
    })

@app.route('/download-summary')
def download_summary():
    file_path = os.path.join(ARCHIVE_FOLDER, "riassunto_referto.docx")
    return send_file(file_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
